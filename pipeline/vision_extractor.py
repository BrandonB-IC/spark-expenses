"""
Claude vision OCR for receipt extraction.

Sends a receipt image or PDF to Claude and parses structured JSON back.

Notable design choices:
- One file -> a LIST of receipts (not a single receipt). PDFs commonly bundle
  multiple receipts (e.g., Uber's monthly statement), so the API always
  returns a list — even if it has only one element.
- Built-in handling for the "initial receipt + tip-updated receipt" pattern
  used by Uber/Lyft/DoorDash. The model is instructed to return only the
  FINAL (tip-updated) version when it sees both for the same trip. This is
  prompt-level dedup; if it proves unreliable in real-world testing, add a
  post-processing pass that compares timestamps + merchant + base amount.
- PDFs are sent natively via the Anthropic API's `document` content block —
  the model sees both the parsed text and a rendered image of each page,
  which is more reliable than splitting to images ourselves.
"""

from __future__ import annotations

import base64
import json
import os
from typing import Optional

from anthropic import Anthropic
from dotenv import load_dotenv

# Make sure ANTHROPIC_API_KEY is loaded from .env when this module is imported
# from a script that hasn't called load_dotenv() itself.
load_dotenv()

_client = Anthropic()
DEFAULT_MODEL = os.getenv("ANTHROPIC_VISION_MODEL", "claude-haiku-4-5-20251001")


EXTRACTION_PROMPT = """You are extracting expense data from a document for an audit-ready reimbursement system.

STEP 1 — CLASSIFY THE DOCUMENT as exactly one of:
- "receipt": proof of a COMPLETED payment — a store/rideshare/hotel/airline
  receipt, an order confirmation showing an amount already paid, or a credit-card
  slip. Most files are receipts. Signals: "paid", a card number, "thank you for
  your order", a transaction/auth code.
- "invoice": a REQUEST for payment or a self-prepared expense summary. It bills
  someone rather than proving payment. Signals: the word "Invoice", an invoice
  number, a "Bill To" / "Billed To" party, "Total Due" / "Amount Due", "Please
  remit payment", or a list of expenses the submitter is asking to be paid for
  WITHOUT attached proof-of-payment receipts.
- "other": anything that is neither (a photo, a note, a blank page).

Put this classification in a "doc_type" field on EVERY row you return.

This file may contain ONE item or MULTIPLE items. Extract every distinct line and
return them as a JSON array.

CRITICAL — INITIAL vs UPDATED RECEIPTS:
Some merchants (especially Uber, Lyft, DoorDash) issue an INITIAL receipt at
the time of service, then issue an UPDATED receipt when the customer adds a
tip. These are NOT separate transactions — they are the SAME transaction with
the tip added. If you see two receipts for the same trip/order at the same or
near-same timestamp with the same merchant and the same base fare, only
include the FINAL (tip-updated) version. Note this in the row's `notes`
field so an auditor can see that a duplicate was suppressed.

CATEGORY is one of these strings (best guess based on merchant + line items):
    "travel-airfare", "travel-hotel", "travel-rideshare", "travel-other",
    "meals", "supplies", "fees", "other"

FOR A RECEIPT, extract these fields per receipt (use null when not visible —
do NOT guess):
- doc_type: "receipt"
- date: ISO 8601 "YYYY-MM-DD" of the transaction
- merchant: business name as printed on the receipt
- merchant_location: city + state if visible (e.g. "San Diego, CA"), else null
- category: see CATEGORY above
- currency: ISO 4217 code (e.g. "USD", "EUR"). Default "USD" if not specified.
- amount: FINAL total paid INCLUDING tax and tip — what hit the card
- subtotal: pre-tax, pre-tip subtotal (null if unclear)
- tax: tax amount (null if not shown)
- tip: tip amount (null if not shown or zero)
- itemization_status: one of:
    "full"      — line items are itemized
    "partial"   — some breakdown but not full line items
    "slip_only" — just a total, e.g. a credit card slip
- line_items: array of {"description": str, "amount": number}.
              Empty array if not itemized.
- notes: short string for any flag worth surfacing to the auditor — e.g.
         "tip-updated version (initial receipt suppressed)",
         "credit card slip — no itemization",
         "foreign currency",
         "duplicate of receipt N suppressed",
         or null if nothing notable.

FOR AN INVOICE, extract ONE row per billed LINE ITEM. Do NOT emit a row for the
invoice SUBTOTAL or TOTAL DUE — only the individual line items (summing the line
rows would otherwise double-count the invoice). Each line-item row:
- doc_type: "invoice"
- date: the date of that line item's expense if shown, else the invoice date
- merchant: the vendor named on that line (e.g. "United Airlines", "Hyatt")
- merchant_location: city + state if visible, else null
- category: see CATEGORY above
- currency: ISO 4217 code. Default "USD".
- amount: the amount billed for that line
- subtotal / tax / tip: null unless the line itself breaks them out
- itemization_status: "invoice_line"
- line_items: [] (or a breakdown if the line itself is itemized)
- invoice_number: the invoice's identifier — the SAME on every line from this invoice
- invoice_note: any billing-arrangement / allocation / cost-sharing note stated
      anywhere on the invoice (e.g. "50% of airfare shared with Boston Retreat
      invoice"), repeated on each line; else null
- notes: anything else notable, else null

Record every extracted line by calling the `record_expense_items` tool with an
"items" array. If nothing is visible, call it with an empty array.
"""


# Forced-tool schema: the model must return its extraction through this tool, so
# we read guaranteed-valid structured data off the tool call instead of parsing
# free text and stripping markdown fences (which occasionally broke). Per-row
# fields stay nullable/loose to mirror the prompt and avoid over-constraining.
EXTRACTION_TOOL = {
    "name": "record_expense_items",
    "description": "Record every distinct receipt line or invoice line item extracted from the document.",
    "input_schema": {
        "type": "object",
        "properties": {
            "items": {
                "type": "array",
                "description": "One entry per receipt or invoice line item. Empty array if nothing is visible.",
                "items": {
                    "type": "object",
                    "properties": {
                        "doc_type": {"type": "string", "enum": ["receipt", "invoice", "other"]},
                        "date": {"type": ["string", "null"]},
                        "merchant": {"type": ["string", "null"]},
                        "merchant_location": {"type": ["string", "null"]},
                        "category": {"type": ["string", "null"]},
                        "currency": {"type": ["string", "null"]},
                        "amount": {"type": ["number", "null"]},
                        "subtotal": {"type": ["number", "null"]},
                        "tax": {"type": ["number", "null"]},
                        "tip": {"type": ["number", "null"]},
                        "itemization_status": {"type": ["string", "null"]},
                        "line_items": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "description": {"type": "string"},
                                    "amount": {"type": "number"},
                                },
                            },
                        },
                        "invoice_number": {"type": ["string", "null"]},
                        "invoice_note": {"type": ["string", "null"]},
                        "notes": {"type": ["string", "null"]},
                    },
                    "required": ["doc_type"],
                },
            }
        },
        "required": ["items"],
    },
}


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _extract_docx_text(file_bytes: bytes) -> str:
    """Pull the text (paragraphs + table cells) out of a .docx file.

    Contractors sometimes submit their expense summary as a Word document rather
    than a receipt. The vision API can't read .docx natively, so we extract the
    text (tables carry the line items) and hand it to the model as text — it then
    classifies it (almost always an 'invoice') and pulls the line items out.
    """
    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(file_bytes))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_msg_text(file_bytes: bytes) -> str:
    """Pull subject + body (+ attachment names) out of an Outlook .msg file.

    Contractors forward email receipts (airfare, hotel, rideshare) as .msg files.
    The vision API can't read .msg, so we parse it with extract-msg and hand the
    text to the model. Note: Drive reports .msg as application/msword, so callers
    must route on the .msg filename extension, not the mime type.
    """
    import os
    import tempfile
    import extract_msg

    tmp = tempfile.NamedTemporaryFile(suffix=".msg", delete=False)
    try:
        tmp.write(file_bytes)
        tmp.close()
        m = extract_msg.openMsg(tmp.name)
        parts = []
        if m.subject:
            parts.append(f"Email subject: {m.subject}")
        if getattr(m, "sender", None):
            parts.append(f"From: {m.sender}")
        if m.body:
            parts.append(m.body)
        atts = [a.longFilename for a in m.attachments if getattr(a, "longFilename", None)]
        if atts:
            parts.append("Attachments: " + ", ".join(atts))
        return "\n".join(parts)
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass


def _build_content_block(file_bytes: bytes, mime_type: str, filename: str = "") -> dict:
    fn = (filename or "").lower()
    if fn.endswith(".msg"):
        text = _extract_msg_text(file_bytes)
        return {
            "type": "text",
            "text": "The following is the text of a forwarded email receipt (.msg file). "
                    "If the subject or body is marked [Personal], it is NOT a business "
                    "expense and should be extracted with a note saying so:\n\n" + text,
        }
    if mime_type == DOCX_MIME or fn.endswith(".docx"):
        text = _extract_docx_text(file_bytes)
        return {
            "type": "text",
            "text": "The following is the extracted text of a submitted .docx expense "
                    "document (tables are rendered as pipe-separated rows):\n\n" + text,
        }
    b64 = base64.standard_b64encode(file_bytes).decode("ascii")
    if mime_type == "application/pdf":
        return {
            "type": "document",
            "source": {
                "type": "base64",
                "media_type": "application/pdf",
                "data": b64,
            },
        }
    if mime_type in ("image/jpeg", "image/png", "image/gif", "image/webp"):
        return {
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": mime_type,
                "data": b64,
            },
        }
    raise ValueError(f"Unsupported mime type for vision extraction: {mime_type}")


def _dedupe_within_file(receipts: list[dict]) -> tuple[list[dict], int]:
    """Drop duplicate receipts that come from the same source file.

    Background: Claude vision sometimes emits a duplicate row for the same
    transaction — e.g., when a multi-receipt PDF has two trips with coincidentally
    identical totals (Uber smoke test, 2026-04-10), the model created a third
    row that was a copy of one of them. This dedup catches those.

    Dedup key: (merchant, date, amount, tip). Extremely unlikely that two
    genuinely distinct receipts share all four values — and if they do, the
    auditor can flag it. False positives are far less harmful than charging
    a contractor twice for the same trip.

    Returns the deduped list and the count of dropped duplicates.
    """
    seen = set()
    out = []
    dropped = 0
    for r in receipts:
        key = (
            (r.get("merchant") or "").strip().lower(),
            r.get("date"),
            r.get("amount"),
            r.get("tip"),
        )
        if key in seen:
            dropped += 1
            continue
        seen.add(key)
        out.append(r)
    return out, dropped


def extract(
    file_bytes: bytes,
    mime_type: str,
    filename: str = "",
    model: Optional[str] = None,
) -> tuple[list[dict], dict]:
    """Run Claude vision on a receipt file and return parsed receipts.

    Returns a tuple of (receipts, metadata) where:
      - receipts is a list of receipt dicts (always a list — even single-receipt files)
      - metadata is {"model", "input_tokens", "output_tokens", "raw_response"}
        for cost tracking and debugging
    """
    content_block = _build_content_block(file_bytes, mime_type, filename)
    use_model = model or DEFAULT_MODEL

    response = _client.messages.create(
        model=use_model,
        max_tokens=4096,
        tools=[EXTRACTION_TOOL],
        tool_choice={"type": "tool", "name": "record_expense_items"},
        messages=[
            {
                "role": "user",
                "content": [
                    content_block,
                    {"type": "text", "text": EXTRACTION_PROMPT},
                ],
            }
        ],
    )

    # Structured output: read the forced tool call — guaranteed schema-valid,
    # no markdown fences to strip.
    data = None
    for block in response.content:
        if getattr(block, "type", None) == "tool_use" and block.name == "record_expense_items":
            data = block.input.get("items", [])
            break
    if data is None:
        raise ValueError(
            f"Claude did not return the expected record_expense_items tool call "
            f"for {filename or '<unnamed>'}."
        )

    if not isinstance(data, list):
        # Defensive: keep the pipeline robust if the tool ever hands back a single object.
        data = [data]

    # Document-type safety default. The classification is per-document, but the
    # model emits it per-row. If ANY row reads as an invoice, treat the WHOLE
    # file as an invoice so its rows are quarantined for manual verification.
    # Biasing toward "invoice" is the safe direction: over-quarantining sends a
    # legit receipt to manual review (annoying), while under-quarantining would
    # auto-reimburse an unsubstantiated invoice (the bug we are preventing).
    has_invoice = any((r.get("doc_type") or "").strip().lower() == "invoice" for r in data)
    file_doc_type = "invoice" if has_invoice else "receipt"
    for r in data:
        raw_dt = (r.get("doc_type") or "receipt").strip().lower()
        r["doc_type"] = "invoice" if has_invoice else (raw_dt if raw_dt in ("receipt", "other") else "receipt")

    # Safety net: drop any duplicate rows the model emitted within this single file.
    data, dropped = _dedupe_within_file(data)

    metadata = {
        "model": use_model,
        "input_tokens": response.usage.input_tokens,
        "output_tokens": response.usage.output_tokens,
        "raw_response": json.dumps({"items": data}),
        "dedup_dropped": dropped,
        "file_doc_type": file_doc_type,
    }
    return data, metadata
